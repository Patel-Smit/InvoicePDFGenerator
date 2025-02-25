""" શ્રી સ્વામિનારાયણાય નમઃ """
import csv
import datetime as dt
import docx
import difflib
import os
from configparser import ConfigParser
from docx2pdf import convert
from tabulate import tabulate

CONSOLE_LENGTH: int = 50
shopping_cart: list = []
services: list = []
services_name: list = []

TAX_RATE: float = 15.00


def main():
    """
    Main entry point of the program. Loads business data from the configuration file,
    initializes the shopping cart and services list, and presents the main menu to the user.

    Calls:
    - `get_cx_details()`
    - `load_available_services()`
    - `main_option_menu()`
    """

    parser = ConfigParser()
    parser.read('./config.ini')
    business_name = parser.get('business_data', 'business_name')

    global shopping_cart, services, services_name
    shopping_cart = []
    services = []
    services_name = []

    print('*' * CONSOLE_LENGTH)
    print(f'*{business_name: ^{CONSOLE_LENGTH - 2}}*')
    print('*' * CONSOLE_LENGTH)
    print(f'{'*' * (CONSOLE_LENGTH - 10)}{' Q. Quit '}*')


    customer_name, customer_address = get_cx_details()
    print(f'\nWelcome, {customer_name} 😎', end='\n')

    load_available_services(services, services_name)
    main_option_menu(customer_name, customer_address, shopping_cart)


def main_option_menu(customer_name, customer_address, cart):
    """
    Displays the main option menu and handles user input. The user can select options to shop, view services,
    view their cart, proceed to checkout, remove items from the cart, or exit the program.

    Args:
    - `customer_name` (str): Name of the customer.
    - `customer_address` (dict): Customer address details.
    - `cart` (list): The shopping cart containing selected services.

    Calls:
    - `_help()`
    - `shop()`
    - `show_available_services()`
    - `show_cart()`
    - `checkout()`
    - `remove_from_cart()`
    """

    while True:
        _help()
        user_input = input('Select your option (h for Help): ').strip()

        match user_input:
            case '1':
                shop(customer_name, customer_address, cart)
            case '2':
                show_available_services(services)
            case '3':
                show_cart(cart)
            case '4':
                checkout(customer_name, customer_address, cart)
            case '5':
                print('')
                if len(cart) == 0:
                    print('🔺Your 🛒 is Empty!')
                    return main_option_menu(customer_name, customer_address, cart)
                remove_from_cart(input('Select the item you want to remove: ').strip().lower(), cart)
            case '6':
                print('')
                print("Thanks for hanging out with us! Until next time, stay awesome! 🚀 Good Bye! 👋🏻")
                print('')
                exit()
            case _:
                print("❌ Invalid Selection!")


def load_available_services(services_list, services_name_list):
    """
    Loads available services from a services_list.csv file and appends them to the provided lists.

    Args:
    - `services_list` (list): A list to store the available services.
    - `services_name_list` (list): A list to store the names of the services.

    Returns:
    - `services_list` (list): The updated services list.
    - `services_name_list` (list): The updated services name list.
    """

    fields = ['service', 'price']
    with open('services_list.csv', 'r', newline='') as file:
        reader = csv.DictReader(file, fieldnames=fields)
        next(reader)
        for row in reader:
            services_list.append(row)
            services_name_list.append(row['service'])
    return services_list, services_name_list


def show_available_services(services_list):
    """
    Displays the list of available services with their prices.

    Args:
    - `services_list` (list): A list of dictionaries containing service data.
    """

    print('')
    print(f'{CONSOLE_LENGTH * '-'}')
    print(f'<{'We provide below services'.center(CONSOLE_LENGTH - 2, '-')}>')
    for service in services_list:
        print(f"{service['service']:30} : ${float(service['price']):.2f}")
    print(f'<{'-' * (CONSOLE_LENGTH - 2)}>')


def get_cx_details():
    """
    Prompts the user to enter their name and address details (unit, street, city, province, and postal code).
    Performs validation on each input to ensure correctness.

    Returns:
    - `cx_name` (str): The customer's name.
    - `cx_address` (dict): The customer's address details.
    """

    while True:
        cx_name = input('Name of Customer: ').strip().title()
        if cx_name.replace(' ', '').isalpha() and cx_name.upper() != 'Q':
            break
        elif cx_name.upper() == 'Q':
            print("Thanks for hanging out with us! Until next time, stay awesome! 🚀 Good Bye! 👋🏻")
            exit()
        else:
            print("⚠️ Please enter a valid name (letters only)")

    print(f'{'Address of Customer':^{CONSOLE_LENGTH}}')

    while True:
        cx_unit = input('Unit/Apt: ').strip().upper()
        if cx_unit.isalnum() or cx_unit == '' and cx_unit.upper() != 'Q':
            break
        elif cx_unit.upper() == 'Q':
            print("Thanks for hanging out with us! Until next time, stay awesome! 🚀 Good Bye! 👋🏻")
            exit()
        else:
            print('⚠️ Please enter a valid unit/apt (Alphanumeric only)')

    while True:
        cx_street = input('Street: ').strip().title()
        if cx_street and cx_street.replace(' ', '').isalnum() and cx_street.upper() != 'Q':
            break
        elif cx_street.upper() == 'Q':
            print("Thanks for hanging out with us! Until next time, stay awesome! 🚀 Good Bye! 👋🏻")
            exit()
        else:
            print('⚠️ Please enter a valid street address (letters and numbers only)')

    while True:
        cx_city_province = input('City, Province: ').strip().title()
        if cx_city_province and cx_city_province.count(',') == 1 and cx_city_province.replace(' ', '').replace(',',
                                                                                                               '').isalpha() and cx_city_province.upper() != 'Q':
            cx_city, cx_province = cx_city_province.split(',')
            cx_city, cx_province = cx_city.strip(), cx_province.strip()
            if not cx_city or not cx_province:
                print('⚠️ Please enter a valid city and province (City, Province)')
                continue
            break
        elif cx_city_province.upper() == 'Q':
            print("Thanks for hanging out with us! Until next time, stay awesome! 🚀 Good Bye! 👋🏻")
            exit()
        elif not cx_city_province.replace(' ', '').replace(',', '').isalpha() and cx_city_province.upper() != 'Q':
            print('⚠️ Please enter a valid city and province')
            print('Format: "City, Province"')
        else:
            print('⚠️ Please enter your city and province (City, Province): ')

    while True:
        cx_postal = input('Postal Code: ').strip().upper()
        if 6 <= len(cx_postal) <= 7 and cx_postal.replace(' ', '').isalnum() and cx_postal.upper() != 'Q':
            break
        elif cx_postal.upper() == 'Q':
            print("Thanks for hanging out with us! Until next time, stay awesome! 🚀 Good Bye! 👋🏻")
            exit()
        else:
            print('⚠️ Please enter your postal code: ')


    if len(cx_province) < 3:
        cx_province = cx_province.upper()

    if cx_unit == '':
        cx_address = {'cx_street': cx_street, 'cx_city': cx_city, 'cx_province': cx_province, 'cx_postal': cx_postal}
    else:
        cx_address = {'cx_unit': cx_unit, 'cx_street': cx_street, 'cx_city': cx_city, 'cx_province': cx_province, 'cx_postal': cx_postal}


    return cx_name, cx_address


def _help():
    """
    Displays the help menu with a list of available options.
    """

    print('')
    print('┏━━━━━━┓')
    print('┃ Menu ┃')
    print('┣━━━━━━┻━━━━━━━━━━━━━━━━━━━━━┓')
    print("┃ 1. Start/Continue Shopping ┃")
    print("┃ 2. List Services Offered   ┃")
    print("┃ 3. View Cart               ┃")
    print("┃ 4. Proceed to Checkout     ┃")
    print("┃ 5. Remove Item from Cart   ┃")
    print("┃ 6. Exit                    ┃")
    print('┗━━━━━━━━━━━━━━━━━━━━━━━━━━━━┛')
    print('')


def shop(cx_name, cx_address, cart):
    """
    Allows the user to select a service and quantity to add to their cart.
    Also provides options to list services or quit to the main menu.

    Args:
    - `cx_name` (str): Name of the customer.
    - `cx_address` (dict): Customer address details.
    - `cart` (list): The shopping cart containing selected services.

    Calls:
    - `show_available_services()`
    - `main_option_menu()`
    - `add_to_cart()`
    - `suggest_item()`
    """

    print('\nL. List offered services')
    print('Q. Quit to the main menu')
    print('Format: "Service xQuantity"\n')
    service = input('Select Service : ').strip()
    match service.upper():
        case 'L':
            show_available_services(services)
            return shop(cx_name, cx_address, cart)
        case 'Q':
            main_option_menu(cx_name, cx_address, cart)
        case '':
            print('\n⚠️ Empty Selection!')
            return shop(cx_name, cx_address, cart)
        case _:
            try:
                service, quantity = service.split(' x')
                quantity = int(quantity)
                if quantity < 0:
                    quantity = 1
            except ValueError:
                quantity = 1

            service_found = False
            for row in services:
                if row['service'].title() == service.title():
                    add_to_cart([row['service'], row['price'], quantity], cart)
                    service_found = True
            if not service_found:
                if suggest_item(service.title(), services_name) == 1:
                    return shop(cx_name, cx_address, cart)
                else:
                    return shop(cx_name, cx_address, cart)

            while True:
                print('')
                add_service = input('Do you want to add service? (Y/N): ').strip().upper()
                if add_service == 'Y':
                    shop(cx_name, cx_address, cart)
                elif add_service == 'N':
                    main_option_menu(cx_name, cx_address, cart)
                else:
                    print('❌ Invalid Input!')


def add_to_cart(item, cart):
    """
    Adds a selected item to the cart or updates its quantity if the item is already present.

    Args:
    - `item` (list): The item to add, containing service name, price, and quantity.
    - `cart` (list): The shopping cart containing selected services.

    Returns:
    - `cart` (list): The updated cart.
    """

    found_index = None

    if len(cart) > 0:
        for index, service_info in enumerate(cart):
            if service_info[0] == item[0]:
                found_index = index
                break

        if found_index is not None:
            cart[found_index][2] += item[2]
            print(f"✅ {cart[found_index][0]} has been updated in cart! Quantity: {cart[found_index][2]}")
            return cart
        else:
            cart.append(item)
            print(f"✅ {item[0]} has been added to cart! Quantity: {item[2]}")
            return cart

    else:
        cart.append(item)
        print(f"✅ {item[0]} has been added to cart! Quantity: {item[2]}")
        return cart


def remove_from_cart(item_name, cart):
    """
    Removes a specified item from the cart or updates its quantity if requested by the user.

    Args:
    - `item_name` (str): The name of the item to remove from the cart.
    - `cart` (list): The shopping cart containing selected services.

    Returns:
    - `cart` (list): The updated cart or 0 if the item is not found.
    """

    if not item_name:
        print(f"❗Empty selection!")
        return 0

    found_index = None

    while True:
        for index, service_info in enumerate(cart):
            if service_info[0].lower() == item_name.lower():
                found_index = index
                break

        if found_index is not None:
            if cart[found_index][2] > 1:
                while True:
                    print(f'You have {cart[found_index][2]} {cart[found_index][0]}')
                    remove_quantity = input('How many you want to remove? : ').strip()
                    try:
                        remove_quantity = int(remove_quantity)
                        if remove_quantity < 0:
                            print('❌ Invalid Quantity! (Only numbers)')
                            continue
                    except ValueError:
                        print('❌ Invalid Quantity! (Only numbers)')
                        continue

                    if remove_quantity == cart[found_index][2]:
                        _item = cart[found_index][0]
                        cart.pop(found_index)
                        print(f'✅ {_item} has been removed from cart')
                        return cart
                    elif remove_quantity < cart[found_index][2]:
                        cart[found_index][2] = cart[found_index][2] - remove_quantity
                        print(f'✅ {cart[found_index][0]} has been updated! Quantity: {cart[found_index][2]}')
                        return cart
                    else:
                        print('❌ Incorrect Quantity!')
            elif cart[found_index][2] == 1:
                _item = cart[found_index][0]
                cart.pop(found_index)
                print(f"✅ {_item} has been removed from cart")
                return cart
            break

        else:
            print(f"❗{item_name} not found in the cart.")
            return 0


def suggest_item(item, items):
    """
    Suggests a similar item name if the user's input does not match any available service.

    Args:
    - `item` (str): The name of the item entered by the user.
    - `items` (list): The list of available item names.

    Returns:
    - `1` if a suggestion is made, `0` if no close match is found.
    """

    suggestions = difflib.get_close_matches(item, items)
    if len(suggestions) == 0:
        print("Item not found!")
        return 0
    else:
        print("Do you mean something like: " + ', '.join(suggestions) + "?")
        return 1


def show_cart(cart):
    """
    Displays the items currently in the cart in a tabular format. Also shows the subtotal of the cart.

    Args:
    - `cart` (list): The shopping cart containing selected services.
    """

    if len(cart) == 0:
        return print('🔺Your 🛒 is Empty!')
    else:
        print('')
        headers = ["Item", "Item Price", "Item Quantity"]
        print(tabulate(cart, headers, tablefmt="heavy_grid"))
        subtotal = 0
        for each_item in cart:
            subtotal = subtotal + float(each_item[1]) * int(each_item[2])

        subtotal = f'{subtotal:.2f}'
        _space = ' ' * (10 - len(subtotal))
        print('┏━━━━━━━━━━━━━━━━━━━━━┓')
        print(f'┃ SubTotal:{_space}{subtotal} ┃')
        print('┗━━━━━━━━━━━━━━━━━━━━━┛')
        print('')


def checkout(cx_name, cx_address, cart):
    """
    Initiates the checkout process by displaying the cart and asking the customer for confirmation.
    If confirmed, proceeds to generate the invoice.

    Args:
    - `cx_name` (str): Name of the customer.
    - `cx_address` (dict): Customer address details.
    - `cart` (list): The shopping cart containing selected services.

    Calls:
    - `generate_invoice()`
    """

    if len(cart) == 0:
        return show_cart(cart)
    else:
        show_cart(cart)
        confirm_checkout = input(f'{cx_name}, confirm checkout? (Y/N): ').strip().upper()
        if confirm_checkout == 'Y':
            return generate_invoice(cx_name, cx_address, cart)
        elif confirm_checkout == 'N':
            return main_option_menu(cx_name, cx_address, cart)
        else:
            print('❌ Invalid Input!')


def replace_text(paragraph, old_text, new_text):
    """
    Replaces old text with new text in a given paragraph.

    Args:
    - `paragraph` (docx.text.paragraph.Paragraph): The paragraph in which to replace text.
    - `old_text` (str): The text to be replaced.
    - `new_text` (str): The new text to replace the old text.
    """

    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(str(old_text), str(new_text))


def generate_invoice(cx_name, cx_address, cart):
    """
    Generates an invoice in Word format and converts it to a PDF. The invoice includes customer and business details,
    a breakdown of the services, and totals (including tax). Updates the invoice number in the configuration file.

    Args:
    - `cx_name` (str): Name of the customer.
    - `cx_address` (dict): Customer address details.
    - `cart` (list): The shopping cart containing selected services.

    Calls:
    - `post_data()`
    """

    parser = ConfigParser()
    parser.read('./config.ini')
    saved_invoice_number = parser.get('settings', 'invoice_number')

    business_name = parser.get('business_data', 'business_name')
    business_street_address = parser.get('business_data', 'business_street_address')
    business_city_province = parser.get('business_data', 'business_city_province')
    business_email_address = parser.get('business_data', 'business_email_address')
    business_contact = parser.get('business_data', 'business_contact')
    business_payment_recipient = parser.get('business_data', 'business_payment_recipient')
    business_payment_bank = parser.get('business_data', 'business_payment_bank')
    business_payment_IBAN = parser.get('business_data', 'business_payment_IBAN')
    business_payment_BIC = parser.get('business_data', 'business_payment_BIC')

    no_of_items = len(cart)
    document = docx.Document('invoice_template.docx')
    row_count = 0
    for i in range(no_of_items, 0, -1):
        table = document.tables[1]

        table.add_row()
        count = 0

        for cells in table.rows[-1].cells:
            if count == 0:
                cells.text = f'[Item{i}]'
                count += 1
            elif count == 1:
                cells.text = f'[Quantity{i}]'
                count += 1
            elif count == 2:
                cells.text = f'[Amount{i}]'
                count += 1
            elif count == 3:
                cells.text = f'[Full Price{i}]'
                count += 1

        insertion_row = table.rows[0]._tr

        insertion_row.addnext(table.rows[-1]._tr)

        if not os.path.exists('temp_files'):
            os.mkdir('temp_files')

        document.save('./temp_files/_sys_modified_template.docx')
        row_count += 1

    doc = docx.Document('./temp_files/_sys_modified_template.docx')

    try:
        customer_name = cx_name
        if 'cx_unit' in cx_address:
            customer_street_address = cx_address['cx_unit'] + ' - ' + cx_address['cx_street']
        else:
            customer_street_address = cx_address['cx_street']
        customer_city = cx_address['cx_city']
        customer_province = cx_address['cx_province']
        customer_postal = cx_address['cx_postal']
        date_time = dt.datetime.today().strftime('%m-%d-%Y') + ' ' + dt.datetime.now().strftime('%I:%M %p')
        replacements = {
            "[Business Name]": business_name,
            "[Business Address]": business_street_address,
            "[Business City Province]": business_city_province,
            "[Business Email]": business_email_address,
            "[Business Contact]": business_contact,
            "[Date]": date_time,
            "[Partner]": customer_name,
            "[Partner Street]": customer_street_address,
            "[Partner City]": customer_city,
            "[Partner Province]": customer_province,
            "[Partner Postal]": customer_postal,
            "[Invoice Number]": saved_invoice_number,
            "[Tax]": f"{TAX_RATE:.2f}",
            "[Recipient]": business_payment_recipient,
            "[Bank]": business_payment_bank,
            "[IBAN]": business_payment_IBAN,
            "[BIC]": business_payment_BIC
        }
        count_ = 1
        current_total = 0
        for item in cart:
            exec(f"replacements['[Item{count_}]'] = '{item[0]}'")
            exec(f"replacements['[Amount{count_}]'] = '$ {item[1]}'")
            exec(f"replacements['[Quantity{count_}]'] = {item[2]}")
            exec(
                f"replacements['[Full Price{count_}]'] = '$ {float(item[1]) * int(item[2]):.2f}'")
            current_total = current_total + float(item[1]) * int(item[2])
            count_ += 1
        count_ = 1
        exec(f"replacements['[Subtotal]'] = '$ {current_total:.2f}'")
        exec(f"replacements['[Total Tax]'] = '$ {current_total * 0.15:.2f}'")
        exec(f"replacements['[Balance Due]'] = '$ {(current_total * 0.15) + current_total:.2f}'")
        total_due = f'$ {(current_total * 0.15) + current_total:.2f}'
    except Exception as e:
        print(e)
        return

    for paragraph in list(doc.paragraphs):
        for old_text, new_text in replacements.items():
            replace_text(paragraph, old_text, new_text)

    for _table in doc.tables:
        for row in _table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        replace_text(paragraph, old_text, new_text)

    doc.save('./temp_files/_system_filled.docx')

    if not os.path.exists('invoices'):
        os.mkdir('invoices')

    convert('./temp_files/_system_filled.docx', f"./invoices/{saved_invoice_number}.pdf")

    post_data(saved_invoice_number, date_time, cx_name, cx_address, cart, total_due)

    set_parser = ConfigParser()
    set_parser.read("./config.ini")
    invoice_number = set_parser['settings']
    invoice_number['invoice_number'] = str(int(saved_invoice_number) + 1).zfill(6)
    with open('./config.ini', 'w') as configfile:
        set_parser.write(configfile)

    print('')
    print('Thank you for shopping with us! 🛍️')
    print('')

    while True:
        user_input = input('Do you want to continue? (Y/N): ').strip().upper()
        if user_input == 'Y':
            cart.clear()
            main()
            break
        elif user_input == 'N':
            print('')
            print("Thanks for hanging out with us! Until next time, stay awesome! 🚀 Good Bye! 👋🏻")
            print('')
            exit()
        else:
            print('❌ Invalid Input!')


def post_data(invoice_number, date_time, cx_name, cx_address, cart, total):
    """
    Saves the sales data (invoice number, date, customer details, cart, and total) to a sales_data.csv file.

    Args:
    - `invoice_number` (str): The generated invoice number.
    - `date_time` (str): The date and time of the transaction.
    - `cx_name` (str): The customer's name.
    - `cx_address` (dict): The customer's address.
    - `cart` (list): The shopping cart containing selected services.
    - `total` (str): The total amount due.
    """

    file_exists = os.path.isfile('sales_data.csv')

    with open('sales_data.csv', 'a', newline='') as f:
        fields = ['invoice_number', 'date_time', 'cx_name', 'cx_address', 'cart', 'total']
        writer = csv.DictWriter(f, fieldnames=fields)

        if not file_exists:
            writer.writeheader()

        writer.writerow({
            'invoice_number': invoice_number,
            'date_time': date_time,
            'cx_name': cx_name,
            'cx_address': cx_address,
            'cart': cart,
            'total': total
        })


if __name__ == "__main__":
    main()
